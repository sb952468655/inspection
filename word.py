# encoding: utf-8
import docx

doc = docx.Document()

# paragraph1 = doc.add_paragraph('hello,world', 'Title')
# paragraph1.add_run(',haha')
# doc.add_paragraph('im wen')

# doc.add_heading('Header 0', 0)
# doc.add_heading('Header 1', 1)
# doc.add_heading('Header 2', 2)
# doc.add_heading('Header 3', 3)
# doc.add_heading('Header 4', 4)

paragraph1 = doc.add_paragraph('hello,world')
paragraph1.runs[0].add_break()
paragraph1.add_run('wwwwww')
doc.add_paragraph('this is a new page')
doc.save('helloworld.docx')

